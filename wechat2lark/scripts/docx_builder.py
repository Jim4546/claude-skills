"""Build a .docx from a parsed WeChat article."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from bs4 import NavigableString, Tag
from docx import Document
from docx.shared import Inches, Pt

from wechat import Article

MAX_IMAGE_WIDTH_INCHES = 6.0
HEADING_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6"}
LIST_TAGS = {"ul", "ol"}
BLOCK_TAGS = HEADING_TAGS | LIST_TAGS | {"p", "blockquote", "pre"}
BOLD_WEIGHTS = {"bold", "bolder", "600", "700", "800", "900"}


def build_docx(article: Article, output_path: Path) -> Path:
    doc = Document()
    doc.add_heading(article.title, level=0)

    meta_parts = []
    if article.author:
        meta_parts.append(article.author)
    if article.publish_date:
        meta_parts.append(article.publish_date.strftime("%Y-%m-%d %H:%M"))
    meta_parts.append(article.url)
    meta_para = doc.add_paragraph(" · ".join(meta_parts))
    for run in meta_para.runs:
        run.font.size = Pt(9)

    _render_children(doc, article.content)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _render_children(doc: Document, node: Tag) -> None:
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                doc.add_paragraph(text)
            continue
        if not isinstance(child, Tag):
            continue
        _render_node(doc, child)


def _render_node(doc: Document, node: Tag) -> None:
    name = node.name.lower() if node.name else ""

    if name == "img":
        _add_image(doc, node)
        return
    if name in HEADING_TAGS:
        level = int(name[1])
        text = node.get_text(strip=True)
        if text:
            doc.add_heading(text, level=min(level, 4))
        return
    if name in LIST_TAGS:
        _render_list(doc, node, ordered=(name == "ol"))
        return
    if name == "blockquote":
        text = node.get_text("\n", strip=True)
        if text:
            doc.add_paragraph(text, style="Intense Quote")
        return
    if name == "pre":
        text = node.get_text("\n", strip=True)
        if text:
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        return
    if name == "hr":
        doc.add_paragraph("───────")
        return
    if name == "table":
        _render_table(doc, node)
        return
    if name == "br":
        return

    # WeChat wraps everything in <section>/<p>/<span> — recurse, but emit a
    # paragraph when we hit a leaf-ish block that has direct text and no block children.
    has_block_child = any(
        isinstance(c, Tag) and c.name and c.name.lower() in BLOCK_TAGS | {"img", "section", "table", "hr"}
        for c in node.children
    )
    if not has_block_child:
        text = node.get_text(" ", strip=True)
        if text:
            level = _heading_level_from_style(_resolve_style(node), text)
            if level is not None:
                try:
                    doc.add_heading(text, level=min(level, 4))
                except Exception:
                    doc.add_paragraph(text)
            else:
                doc.add_paragraph(text)
        # still pick up nested images that markup separately
        for img in node.find_all("img", recursive=True):
            _add_image(doc, img)
        return

    _render_children(doc, node)


def _render_list(doc: Document, node: Tag, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    for li in node.find_all("li", recursive=False):
        text = li.get_text(" ", strip=True)
        if text:
            try:
                doc.add_paragraph(text, style=style)
            except KeyError:
                doc.add_paragraph(("- " if not ordered else "1. ") + text)
        for img in li.find_all("img"):
            _add_image(doc, img)


def _render_table(doc: Document, node: Tag) -> None:
    rows = node.find_all("tr")
    if not rows:
        return
    cols = max(len(r.find_all(["td", "th"])) for r in rows)
    if cols == 0:
        return
    table = doc.add_table(rows=len(rows), cols=cols)
    for i, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        for j, cell in enumerate(cells):
            table.cell(i, j).text = cell.get_text(" ", strip=True)


_STYLE_DECL_RE = re.compile(r"\s*([A-Za-z\-]+)\s*:\s*([^;]+?)\s*(?:;|$)")
_PX_RE = re.compile(r"(\d+(?:\.\d+)?)\s*px", re.IGNORECASE)


def _parse_inline_style(text: str) -> dict[str, str]:
    out: dict[str, str] = {}
    for m in _STYLE_DECL_RE.finditer(text or ""):
        out[m.group(1).strip().lower()] = m.group(2).strip().lower()
    return out


def _resolve_style(node: Tag) -> dict[str, str]:
    """Merge inline styles up the parent chain. Leaf wins on conflict."""
    chain: list[dict[str, str]] = []
    cur = node
    while cur is not None and isinstance(cur, Tag):
        style_attr = cur.attrs.get("style") if cur.attrs else None
        if style_attr:
            chain.append(_parse_inline_style(style_attr))
        cur = cur.parent
    merged: dict[str, str] = {}
    for s in reversed(chain):
        merged.update(s)
    return merged


def _heading_level_from_style(style: dict[str, str], text: str) -> Optional[int]:
    """Map inline style to a heading level (1-4) or None for body text."""
    size = _PX_RE.search(style.get("font-size", ""))
    if not size:
        return None
    size_px = float(size.group(1))
    weight = style.get("font-weight", "400").strip()
    is_bold = weight in BOLD_WEIGHTS
    text_len = len(text)

    if size_px >= 24 and text_len <= 60:
        return 1
    if size_px >= 19 and text_len <= 80:
        return 2
    if size_px >= 16 and is_bold and text_len <= 60:
        return 3
    if size_px >= 14 and is_bold and text_len <= 40:
        return 4
    return None


def _add_image(doc: Document, img: Tag) -> None:
    local = img.get("data-local")
    if not local:
        return
    path = Path(local)
    if not path.exists():
        return
    try:
        doc.add_picture(str(path), width=Inches(MAX_IMAGE_WIDTH_INCHES))
    except Exception:
        # Some images (e.g. webp with quirks) may fail; skip rather than abort the whole doc.
        doc.add_paragraph(f"[image skipped: {path.name}]")
