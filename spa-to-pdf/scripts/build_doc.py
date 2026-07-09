"""Build a Word document from raw/index.json.

Usage: python build_doc.py [--in raw/index.json] [--out output.docx]

The input JSON is what scrape.py produces:
    {root_url, page_title, subtitle, sections: [{title, html, links}, ...]}

The output docx contains:
  - Cover (page_title + subtitle + source link)
  - Bilingual TOC field (auto-fills when Word opens or via to_pdf.py)
  - One Heading 1 per section (page-break-before, no trailing blank pages)
  - Original HTML rendered: paragraphs, nested lists, bold/italic, hyperlinks

See ../references/pitfalls.md for the design decisions behind each quirk.
"""
import argparse
import json
import re
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# Strip XML-illegal control chars (allow \t \n \r)
_CTRL_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]")


def clean(s: str) -> str:
    return _CTRL_RE.sub("", s) if s else ""


def add_hyperlink(paragraph, url: str, text: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = clean(text or url)
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_toc(doc):
    """Insert a Word TOC field. to_pdf.py updates it; manually in Word press F9."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_el = run._r
    r_el.append(fld_begin)
    r_el.append(instr)
    r_el.append(fld_sep)
    r_el.append(fld_end)


def set_run_fonts(run, ascii_font="Calibri", eastasia_font="Microsoft YaHei"):
    """Explicitly set Latin and East-Asian fonts for visual consistency."""
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)
    rFonts.set(qn("w:eastAsia"), eastasia_font)
    rFonts.set(qn("w:cs"), ascii_font)
    rPr.append(rFonts)


def render_inline(paragraph, node, bold=False, italic=False):
    if isinstance(node, NavigableString):
        txt = clean(str(node))
        if not txt:
            return
        run = paragraph.add_run(txt)
        run.bold = bold
        run.italic = italic
        return
    if not isinstance(node, Tag):
        return

    name = node.name.lower()
    if name == "a":
        href = clean(node.get("href", "").strip())
        text = clean(node.get_text(" ", strip=True)) or href
        if href:
            add_hyperlink(paragraph, href, text)
        else:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
        return
    if name == "br":
        paragraph.add_run().add_break()
        return
    new_bold = bold or name in ("b", "strong")
    new_italic = italic or name in ("i", "em")
    for child in node.children:
        render_inline(paragraph, child, new_bold, new_italic)


def _list_style(tag_name: str, level: int) -> str:
    base = "List Bullet" if tag_name == "ul" else "List Number"
    return base if level <= 1 else f"{base} {min(level, 3)}"


def _render_list(doc, list_node, level: int):
    """Render <ul>/<ol>, tolerating non-standard nested <ul>/<ol> as direct
    siblings of <li> (an HTML quirk seen in some SPAs)."""
    style = _list_style(list_node.name.lower(), level)
    for child in list_node.children:
        if not isinstance(child, Tag):
            continue
        cname = child.name.lower()
        if cname == "li":
            p = doc.add_paragraph(style=style)
            for sub in child.children:
                if isinstance(sub, Tag) and sub.name.lower() in ("ul", "ol"):
                    _render_list(doc, sub, level + 1)
                else:
                    render_inline(p, sub)
        elif cname in ("ul", "ol"):
            _render_list(doc, child, level + 1)


def render_block(doc, node):
    if isinstance(node, NavigableString):
        text = str(node).strip()
        if text:
            doc.add_paragraph().add_run(clean(text))
        return
    if not isinstance(node, Tag):
        return

    name = node.name.lower()
    if name in ("script", "style"):
        return

    if name in ("p", "div"):
        block_children = [c for c in node.children if isinstance(c, Tag) and c.name in (
            "p", "div", "ul", "ol", "h1", "h2", "h3", "h4", "h5", "h6", "blockquote", "hr")]
        if block_children:
            for c in node.children:
                render_block(doc, c)
            return
        text = node.get_text(strip=True)
        if not text:
            return
        p = doc.add_paragraph()
        render_inline(p, node)
        return

    if name in ("ul", "ol"):
        _render_list(doc, node, level=1)
        return

    if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(name[1])
        doc.add_heading(clean(node.get_text(" ", strip=True)), level=min(level + 2, 4))
        return

    if name == "blockquote":
        p = doc.add_paragraph(style="Intense Quote")
        render_inline(p, node)
        return

    if name == "hr":
        doc.add_paragraph("―" * 20)
        return

    text = node.get_text(strip=True)
    if text:
        p = doc.add_paragraph()
        render_inline(p, node)


def add_section_html(doc, html: str):
    soup = BeautifulSoup(html, "lxml")
    root = soup.body or soup
    # If we wrapped a content div, descend once into it
    if root.name == "html":
        root = root.body or root
    for child in root.children:
        render_block(doc, child)


def build(data: dict, out_path: Path):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Cover
    title = doc.add_paragraph()
    run = title.add_run(clean(data.get("page_title", "")))
    run.bold = True
    run.font.size = Pt(28)
    set_run_fonts(run)
    if data.get("subtitle"):
        sub = doc.add_paragraph()
        sub_run = sub.add_run(clean(data["subtitle"]))
        sub_run.italic = True
        sub_run.font.size = Pt(16)
        set_run_fonts(sub_run)
    if data.get("root_url"):
        src = doc.add_paragraph()
        src_run = src.add_run("Source: ")
        src_run.font.size = Pt(9)
        add_hyperlink(src, data["root_url"], data["root_url"])
    doc.add_page_break()

    # TOC (plain paragraph as title, NOT Heading 1, to avoid self-reference)
    toc_title = doc.add_paragraph()
    tr = toc_title.add_run("目录 / Table of Contents")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = RGBColor(0x2E, 0x55, 0x95)
    set_run_fonts(tr)
    add_toc(doc)
    doc.add_page_break()

    # Sections (page-break-before on H1, no trailing add_page_break)
    for i, sec in enumerate(data["sections"]):
        h = doc.add_heading(clean(sec["title"]), level=1)
        if i > 0:
            h.paragraph_format.page_break_before = True
        add_section_html(doc, sec["html"])

    doc.save(out_path)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--in", dest="inp", default="raw/index.json")
    ap.add_argument("--out", default="output.docx")
    args = ap.parse_args()

    data = json.loads(Path(args.inp).read_text(encoding="utf-8"))
    out = Path(args.out)
    build(data, out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
